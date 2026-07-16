from pathlib import Path

from docx import Document as DocxDocument
import pytest
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from app.database import Base
from app.document_service import extract_sections, parse_and_rechunk, save_document_version
from app.models import ContentVersion, Document, DocumentChunk
from app.routers import documents
from app.schemas import DocumentChunkUpdate, DocumentUpdate


@pytest.fixture()
def session():
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    with Session(engine) as db:
        yield db
    Base.metadata.drop_all(engine)


def make_document(session: Session, filename: str = "rag-notes.md") -> Document:
    document = Document(
        title="RAG Notes",
        slug="rag-notes",
        original_filename=filename,
        stored_filename=filename,
        content_type="text/markdown",
        visibility="private",
        status="processing",
        chunk_size=200,
        chunk_overlap=20,
    )
    session.add(document)
    session.flush()
    return document


def test_markdown_parse_chunk_edit_and_version_restore(session, tmp_path: Path):
    source = tmp_path / "rag-notes.md"
    source.write_text(
        "# Hybrid Search\n\nMilvus vector search combines with BM25 lexical search.\n\n"
        "## Reranking\n\nA reranker improves the final evidence order.",
        encoding="utf-8",
    )
    document = make_document(session)
    assert parse_and_rechunk(session, document, source) == 2
    save_document_version(session, document, "admin@example.com", "parsed")
    session.commit()

    chunks = list(session.scalars(select(DocumentChunk).order_by(DocumentChunk.chunk_index)))
    assert [chunk.heading for chunk in chunks] == ["Hybrid Search", "Reranking"]
    assert chunks[0].embedding_dimensions >= 16
    assert document.status == "ready"
    assert "Milvus vector search" in document.raw_text

    initial_version = session.scalar(select(ContentVersion).where(ContentVersion.reason == "parsed"))
    updated = documents.update_document_chunk(
        chunks[0].id,
        DocumentChunkUpdate(
            heading="Hybrid Retrieval",
            content="Dense retrieval and BM25 are fused before reranking.",
            metadata={"reviewed": True},
            is_enabled=False,
        ),
        user="admin@example.com",
        session=session,
    )
    assert updated["heading"] == "Hybrid Retrieval"
    assert updated["is_enabled"] is False

    restored = documents.restore_document_version(
        initial_version.id,
        user="admin@example.com",
        session=session,
    )
    assert restored["chunks"][0]["heading"] == "Hybrid Search"
    assert restored["chunks"][0]["is_enabled"] is True
    assert restored["revision"] == 3


def test_document_metadata_update_records_version(session):
    document = make_document(session)
    document.status = "ready"
    session.commit()

    result = documents.update_document(
        document.id,
        DocumentUpdate(
            title="Enterprise RAG Notes",
            slug="enterprise-rag-notes",
            summary="Document knowledge base notes.",
            visibility="unlisted",
            allow_ai_search=False,
            metadata={"source": "internal"},
            expected_revision=1,
        ),
        user="admin@example.com",
        session=session,
    )
    assert result["revision"] == 2
    assert result["visibility"] == "unlisted"
    assert result["allow_ai_search"] is False
    assert session.scalar(select(ContentVersion).where(ContentVersion.reason == "manual_save")) is not None


def test_docx_parser_preserves_heading_sections(tmp_path: Path):
    source = tmp_path / "agent-memory.docx"
    docx = DocxDocument()
    docx.add_heading("Agent Memory", level=1)
    docx.add_paragraph("Memory stores durable facts for later agent runs.")
    docx.add_heading("Retrieval", level=2)
    docx.add_paragraph("Relevant memories are retrieved before planning.")
    docx.save(source)

    sections = extract_sections(source)
    assert [section.heading for section in sections] == ["Agent Memory", "Retrieval"]
    assert "durable facts" in sections[0].text
